# resume_logic.py
import fitz  # PyMuPDF for PDF
from docx import Document
import re
import tempfile
from docx2pdf import convert

def extract_text_from_resume(file_path: str) -> str:
    """Extract text from docx or pdf resume."""
    text = ""
    if file_path.endswith(".docx"):
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif file_path.endswith(".pdf"):
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text("text")
    return text

def extract_skills_from_text(text: str) -> list:
    """Very basic skill extraction (can be improved with NLP)."""
    skills_db = ["Python", "SQL", "Power BI", "Tableau", "Machine Learning", 
                 "Deep Learning", "NLP", "Excel", "FastAPI", "GCP", "AWS", "Agile"]
    found = [skill for skill in skills_db if re.search(rf"\b{skill}\b", text, re.IGNORECASE)]
    return found

def compare_skills(resume_skills: list, jd_skills: list) -> dict:
    """Return matched and missing skills."""
    missing = [s for s in jd_skills if s not in resume_skills]
    matched = [s for s in jd_skills if s in resume_skills]
    return {"matched": matched, "missing": missing}

def create_custom_resume(base_resume: str, jd_text: str, approved_skills: list, company_name: str) -> str:
    """Insert approved skills into resume and generate PDF."""
    doc = Document(base_resume)

    # Find placeholder in resume
    for para in doc.paragraphs:
        if "{{skills_placeholder}}" in para.text:
            para.text = para.text.replace("{{skills_placeholder}}", ", ".join(approved_skills))

    # Save as new docx
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_docx)

    # Convert to PDF
    output_pdf = f"CV_{company_name}.pdf"
    convert(temp_docx, output_pdf)

    return output_pdf
