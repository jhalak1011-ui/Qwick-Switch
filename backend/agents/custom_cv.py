import os
import re
import json
import PyPDF2
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import anthropic
from dotenv import load_dotenv

load_dotenv("openaiapikey.env")
_client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))


# ─── Text extraction ───────────────────────────────────────────────────────────

def extract_resume_text(path: str) -> str:
    if not path or not os.path.exists(path):
        return ""
    try:
        if path.lower().endswith(".pdf"):
            text = ""
            with open(path, "rb") as f:
                for page in PyPDF2.PdfReader(f).pages:
                    text += (page.extract_text() or "") + "\n"
            return text.strip()
        if path.lower().endswith(".docx"):
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception:
        pass
    return ""


# ─── ATS score ─────────────────────────────────────────────────────────────────

def calculate_ats_score(resume_text: str, keywords: list) -> int:
    if not keywords or not resume_text:
        return 0
    text_lower = resume_text.lower()
    found = sum(1 for kw in keywords if kw.lower() in text_lower)
    return round((found / len(keywords)) * 100)


# ─── Keyword extraction ────────────────────────────────────────────────────────

_KW_SYSTEM = """\
You are an ATS and CV expert. Analyze the job description and extract keywords.

Return ONLY a JSON object:
{
  "required": [{"keyword": "...", "frequency": N, "in_resume": true/false}],
  "preferred": [{"keyword": "...", "frequency": N, "in_resume": true/false}],
  "responsibilities": [{"keyword": "...", "frequency": N, "in_resume": true/false}],
  "industry_terms":   [{"keyword": "...", "frequency": N, "in_resume": true/false}]
}

Rules:
- required: explicitly required/mandatory skills and qualifications
- preferred: nice-to-have or bonus skills
- responsibilities: key tasks/action phrases from the role (e.g. "cross-functional collaboration")
- industry_terms: domain tools and terminology (e.g. "agile", "Salesforce")
- frequency: number of times keyword appears in the JD
- in_resume: true if keyword or close synonym already exists in the resume
- 5–8 items per category max, most important first
- Return ONLY raw JSON, no markdown, no explanation\
"""

def extract_keywords(jd_text: str, resume_text: str) -> dict:
    user = f"Job Description:\n{jd_text[:3000]}\n\nResume:\n{resume_text[:2000]}"
    try:
        resp = _client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1024,
            system=_KW_SYSTEM,
            messages=[{"role": "user", "content": user}],
        )
        raw = resp.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.I).strip()
        return json.loads(raw)
    except Exception as e:
        return {"required": [], "preferred": [], "responsibilities": [], "industry_terms": [], "error": str(e)}


# ─── CV tailoring ──────────────────────────────────────────────────────────────

_TAILOR_SYSTEM = """\
You are an expert CV writer. Tailor the existing CV to a job description using only the provided keywords.

STRICT RULES — violating these makes the output useless:
- Do NOT add experience, companies, degrees, or skills not already in the CV
- Only REPHRASE existing bullets to naturally surface the selected keywords
- Keep all facts: dates, company names, job titles, figures, numbers
- Rewrite the professional summary to mirror the JD tone
- Each rewritten bullet must be clearly traceable to an original bullet

Return ONLY a JSON object:
{
  "summary":    {"original": "...", "rewritten": "..."},
  "experience": [{"original": "...", "rewritten": "..."}],
  "skills":     {"original": "...", "rewritten": "..."},
  "keywords_added": ["kw1", "kw2"]
}

- Omit sections not present in the CV
- "experience" is a list of bullet-level diffs (one object per bullet point)
- Return ONLY raw JSON, no markdown\
"""

def tailor_cv(resume_text: str, selected_keywords: list, jd_text: str) -> dict:
    user = (
        f"Keywords to incorporate: {', '.join(selected_keywords)}\n\n"
        f"Job Description:\n{jd_text[:1500]}\n\n"
        f"Original CV:\n{resume_text[:3500]}"
    )
    try:
        resp = _client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2500,
            system=_TAILOR_SYSTEM,
            messages=[{"role": "user", "content": user}],
        )
        raw = resp.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.I).strip()
        return json.loads(raw)
    except Exception as e:
        return {"error": str(e)}


# ─── DOCX in-place editor ──────────────────────────────────────────────────────

def _normalize(text: str) -> str:
    """Normalize whitespace for fuzzy paragraph matching."""
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _replace_para_text(para, new_text: str):
    """
    Replace a paragraph's text while preserving run-level formatting
    (bold, italic, font size, colour, etc.).
    Puts all new text into the first run, clears the rest.
    """
    if not para.runs:
        # paragraph has no runs — add one
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def build_cv_docx(diffs: dict, approved: dict, profile: dict, output_path: str):
    """
    Edit the user's original DOCX CV in-place:
    - Find each original bullet by text matching
    - Replace its text with the approved rewrite
    - Save to output_path, preserving all original formatting

    Falls back to a plain text DOCX if no original DOCX is available.
    """
    original_docx = profile.get("resume_docx_file", "")

    if not original_docx or not os.path.exists(original_docx):
        _build_fallback_docx(diffs, approved, profile, output_path)
        return

    doc = Document(original_docx)

    # Build change map: normalized_original_text → new_text
    change_map: dict[str, str] = {}

    summary = diffs.get("summary", {})
    if summary.get("original"):
        change_map[_normalize(summary["original"])] = (
            approved.get("summary") or summary.get("rewritten") or summary["original"]
        )

    for i, bullet in enumerate(diffs.get("experience", [])):
        orig = bullet.get("original", "")
        if orig:
            change_map[_normalize(orig)] = (
                approved.get(f"exp_{i}") or bullet.get("rewritten") or orig
            )

    skills = diffs.get("skills", {})
    if skills.get("original"):
        change_map[_normalize(skills["original"])] = (
            approved.get("skills") or skills.get("rewritten") or skills["original"]
        )

    def _process_paragraphs(paragraphs):
        for para in paragraphs:
            norm = _normalize(para.text)
            if norm in change_map:
                _replace_para_text(para, change_map[norm])

    # Main body paragraphs
    _process_paragraphs(doc.paragraphs)

    # Paragraphs inside tables (some CVs use table-based layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(cell.paragraphs)

    doc.save(output_path)


def _build_fallback_docx(diffs: dict, approved: dict, profile: dict, output_path: str):
    """Plain fallback DOCX when no original DOCX is available."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.7)
        sec.left_margin = sec.right_margin = Inches(0.85)

    name = profile.get("name", "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name); r.bold = True; r.font.size = Pt(16)

    contact = " | ".join(x for x in [
        profile.get("email",""), profile.get("phone","") or "",
        profile.get("linkedin","") or ""
    ] if x)
    if contact:
        cp = doc.add_paragraph(contact)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = diffs.get("summary", {})
    if summary:
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(approved.get("summary") or summary.get("rewritten",""))

    if diffs.get("experience"):
        doc.add_heading("Experience", level=2)
        for i, b in enumerate(diffs["experience"]):
            doc.add_paragraph(
                approved.get(f"exp_{i}") or b.get("rewritten",""),
                style="List Bullet"
            )

    if diffs.get("skills"):
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(
            approved.get("skills") or diffs["skills"].get("rewritten","")
        )

    doc.save(output_path)
